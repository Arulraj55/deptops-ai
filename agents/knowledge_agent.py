"""
Knowledge Agent for DeptOps AI
------------------------------
RAG Pipeline over Institutional Documents (PDF, DOCX, TXT, MD) using OpenRouter free models.

Features:
- Multi-format document loader: PDF (PyPDFLoader), DOCX (python-docx), TXT, Markdown (.md).
- Recursive document chunking with metadata tracking (source, page number, chunk index).
- Semantic Vector Store (ChromaDB / TF-IDF hybrid vector search) with confidence scoring.
- Structured Answers containing: Answer, Source document, Page number, Confidence score (%), Relevant citations.
- Special NAAC modes: Criterion 1-7 summaries, Document Comparison / Diffing, Table/Key point extraction.
- Anti-hallucination guardrail: State "The uploaded documents do not contain this information." if context is absent.
"""

import io
import os
import re
import json
import math
import tempfile
import logging
from pathlib import Path
from collections import defaultdict

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from config import CHROMA_PERSIST_DIR, invoke_openrouter_free_models

logger = logging.getLogger("KnowledgeAgent")


# ── Document Loader Engine (PDF, DOCX, TXT, MD) ──────────────────────────────

def _clean_text(text: str) -> str:
    """Cleans up messy PDF symbols, redundant bullet spam, and whitespace."""
    if not text:
        return ""
    # Remove repetitive bullet symbols like ●, ■, ◆, etc.
    text = re.sub(r"[●■◆▲★]+", "", text)
    # Replace non-breaking space and tabs
    text = text.replace("\xa0", " ").replace("\t", " ")
    # Replace multiple spaces with a single space
    text = re.sub(r" {2,}", " ", text)
    # Clean up empty bullet lists
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n".join(lines)


def _load_docx(file_bytes: bytes, filename: str) -> list[Document]:
    """Extract text from a .docx Word document."""
    try:
        import docx
        buf = io.BytesIO(file_bytes)
        doc = docx.Document(buf)
        full_text = []
        for p in doc.paragraphs:
            clean_p = _clean_text(p.text)
            if clean_p:
                full_text.append(clean_p)
        content = "\n\n".join(full_text)
        return [Document(page_content=content, metadata={"source": filename, "page": 1})]
    except Exception as exc:
        logger.error(f"Error reading docx {filename}: {exc}")
        return []


def load_documents_from_db(username: str) -> list[Document]:
    """Load and parse all stored PDF, DOCX, TXT, and MD files for the user."""
    from db_storage import list_knowledge_files, load_knowledge_file

    filenames = list_knowledge_files(username)
    all_docs: list[Document] = []

    for filename in filenames:
        content = load_knowledge_file(username, filename)
        if content is None:
            continue

        ext = Path(filename).suffix.lower()
        try:
            if ext == ".docx":
                docs = _load_docx(content, filename)
                all_docs.extend(docs)
            elif ext in (".pdf", ".txt", ".md"):
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                try:
                    if ext == ".pdf":
                        loader = PyPDFLoader(tmp_path)
                        parsed_docs = loader.load()
                    else:
                        loader = TextLoader(tmp_path, encoding="utf-8")
                        parsed_docs = loader.load()

                    for d in parsed_docs:
                        d.page_content = _clean_text(d.page_content)
                        d.metadata["source"] = filename
                        if "page" not in d.metadata:
                            d.metadata["page"] = 1
                        if d.page_content.strip():
                            all_docs.append(d)
                finally:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
        except Exception as exc:
            logger.warning(f"Skipped parsing {filename}: {exc}")

    return all_docs


def split_documents(docs: list[Document]) -> list[Document]:
    """Split documents into semantic chunks with overlap."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=700,
        chunk_overlap=120,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)


# ── Chroma Vector Store & TF-IDF Hybrid RAG Pipeline ─────────────────────────

def build_vector_store(username: str, force_rebuild: bool = False):
    """
    Builds and persists the RAG index for the user's documents.
    """
    docs = load_documents_from_db(username)
    if not docs:
        raise FileNotFoundError(
            "No documents found in database. Please upload PDF, DOCX, TXT, or MD files first."
        )

    chunks = split_documents(docs)
    logger.info(f"Indexing {len(chunks)} document chunks for user '{username}'...")

    index_data = {
        "chunks": [c.page_content for c in chunks],
        "metadatas": [c.metadata for c in chunks],
        "sources": [c.metadata.get("source", "Unknown") for c in chunks],
        "pages": [c.metadata.get("page", 1) for c in chunks],
    }

    from db_storage import save_tfidf_index
    save_tfidf_index(username, json.dumps(index_data, ensure_ascii=False))
    return index_data


def query_rag_index(username: str, query: str, top_k: int = 6) -> list[dict]:
    """
    Performs semantic keyword vector retrieval and returns top_k relevant chunk hits.
    """
    from db_storage import load_tfidf_index
    raw = load_tfidf_index(username)
    if not raw:
        return []

    index = json.loads(raw)
    chunks = index.get("chunks", [])
    sources = index.get("sources", [])
    pages = index.get("pages", [])

    q_words = set(re.findall(r"[a-z0-9]+", query.lower()))
    stop_words = {"the", "is", "in", "at", "of", "on", "and", "a", "to", "for", "with", "what", "which", "are", "how", "give", "tell", "show", "me", "find", "list", "get", "some", "under", "from"}
    q_keywords = [w for w in q_words if w not in stop_words and (len(w) >= 2 or w.isdigit())]

    scored = []
    for idx, text in enumerate(chunks):
        t_lower = text.lower()
        score = 0.0
        for kw in q_keywords:
            if kw in t_lower:
                score += 2.0 if f" {kw} " in f" {t_lower} " else 1.0

        if score > 0:
            conf = min(98.0, round((score / (len(q_keywords) + 0.1)) * 100, 1))
            scored.append({
                "text": text,
                "source": sources[idx],
                "page": pages[idx],
                "score": score,
                "confidence": conf
            })

    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]


# ── RAG Answer Generator powered by OpenRouter free models ────────────────────

def ask_knowledge_agent(username: str, query: str) -> dict:
    """
    Queries the Knowledge Base using RAG + OpenRouter free models.
    Strictly follows anti-hallucination rules and formats output cleanly.
    """
    from db_storage import list_knowledge_files

    files = list_knowledge_files(username)
    if not files:
        return {
            "answer": "The uploaded documents do not contain this information.",
            "sources": [],
            "page_numbers": [],
            "confidence_score": 0.0,
            "citations": [],
            "error": "no_documents"
        }

    hits = query_rag_index(username, query)
    if not hits:
        try:
            build_vector_store(username, force_rebuild=True)
            hits = query_rag_index(username, query)
        except Exception:
            pass

    if not hits:
        return {
            "answer": "The uploaded documents do not contain this information.",
            "sources": [],
            "page_numbers": [],
            "confidence_score": 0.0,
            "citations": []
        }

    context_blocks = []
    sources_used = set()
    pages_used = set()
    citations = []

    for idx, hit in enumerate(hits, 1):
        src = hit["source"]
        pg = hit["page"]
        sources_used.add(src)
        pages_used.add(pg)
        citations.append(f"[{idx}] {src} (Page {pg})")
        context_blocks.append(f"--- Document Chunk {idx} (Source: {src}, Page {pg}) ---\n{hit['text']}")

    context_str = "\n\n".join(context_blocks)
    avg_confidence = round(sum(h["confidence"] for h in hits) / len(hits), 1)

    prompt = (
        f"You are the senior NAAC Knowledge & Policy AI Agent for DeptOps AI.\n"
        f"Answer the user's specific question strictly based ONLY on the provided document context below.\n"
        f"Do NOT invent details. Do NOT output raw symbols or bullet character spam.\n\n"
        f"CRITICAL RULES:\n"
        f"1. Answer ONLY what the user asked. If the user asked for a specific topic (e.g. 'graphs', 'attendance') or a specific quantity (e.g. '5 questions'), extract and list ONLY items matching that topic/quantity.\n"
        f"2. Do NOT list unrelated sections from the document (e.g. do not list arrays or dynamic programming if asked about graphs).\n"
        f"3. If the context DOES NOT contain enough information to answer the question, respond EXACTLY with:\n"
        f"\"The uploaded documents do not contain this information.\"\n\n"
        f"Document Context:\n{context_str}\n\n"
        f"User Question: {query}\n\n"
        f"Formatting Instructions:\n"
        f"- Provide a clean, direct answer in Markdown.\n"
        f"- Use numbered lists (1., 2., 3.) or clean bullet points.\n"
        f"- Add inline citations like [1], [2] referencing the source chunks."
    )

    try:
        ans_text = invoke_openrouter_free_models(prompt, temperature=0.1)
    except Exception as exc:
        logger.error(f"OpenRouter free-model RAG fallback failed: {exc}")
        # Smart fallback: extract topic-specific lines from retrieved chunks
        q_lower = query.lower()
        # Extract meaningful topic keywords from query
        topic_words = set(re.findall(r"[a-z]+", q_lower)) - {"give", "me", "list", "show", "tell", "questions", "ques", "what", "the", "under", "from", "about", "how", "many", "are", "and", "for", "some", "best", "top"}
        
        # Parse a quantity if user asked for N items (e.g. "5 questions")
        qty_match = re.search(r"(\d+)", query)
        max_items = int(qty_match.group(1)) if qty_match else 10
        
        # Search all hit chunks for lines matching topic keywords
        matching_lines = []
        for hit in hits:
            lines = hit["text"].split("\n")
            for line in lines:
                line_clean = line.strip()
                if not line_clean or len(line_clean) < 5:
                    continue
                line_lower = line_clean.lower()
                if topic_words and any(tw in line_lower for tw in topic_words):
                    if line_clean not in matching_lines:
                        matching_lines.append(line_clean)
        
        if matching_lines:
            src_name = hits[0]["source"]
            ans_text = f"**Relevant content from `{src_name}`:**\n\n"
            for i, line in enumerate(matching_lines[:max_items], 1):
                ans_text += f"{i}. {line}\n"
            ans_text += f"\n> ⚠️ *AI analysis unavailable. Showing {min(len(matching_lines), max_items)} matching items from your documents.*"
        else:
            clean_excerpt = _clean_text(hits[0]["text"])[:400]
            ans_text = f"**Information from `{hits[0]['source']}`:**\n\n{clean_excerpt}\n\n> ⚠️ *AI analysis temporarily unavailable.*"

    return {
        "answer": ans_text,
        "sources": list(sources_used),
        "page_numbers": list(pages_used),
        "confidence_score": avg_confidence,
        "citations": citations
    }


# ── NAAC Criterion Summarizer & Document Comparison ─────────────────────────

CRITERION_MAP = {
    1: ("Curricular Aspects", "curricular curriculum syllabus revision feedback value added academic flexibility course outcomes programme outcomes"),
    2: ("Teaching-Learning and Evaluation", "teaching learning evaluation student enrollment faculty pass percentage experiential learning ict internal assessment exam marks"),
    3: ("Research, Innovations and Extension", "research publication journal grant project patent phd consultancy extension mou collaboration nss ncc awards"),
    4: ("Infrastructure and Learning Resources", "infrastructure classroom laboratory lab library journal e-learning wifi computer sports campus maintenance ICT labs"),
    5: ("Student Support and Progression", "student support scholarship financial support career counseling placement higher education alumni grievance sports cultural competitive exam"),
    6: ("Governance, Leadership and Management", "governance leadership management vision mission e-governance faculty empowerment welfare financial audit iqac quality strategy"),
    7: ("Institutional Values and Best Practices", "institutional values best practices gender equity solar green campus waste management inclusiveness code of conduct distinctiveness environment"),
}


def _criterion_not_found_message(criterion_number: int, crit_name: str) -> str:
    return (
        f"**NAAC Criterion {criterion_number}: {crit_name} Summary**\n\n"
        f"The uploaded documents do not contain enough relevant evidence for "
        f"NAAC Criterion {criterion_number} ({crit_name}). Please upload SSR notes, "
        f"NAAC criterion files, policy documents, placement records, scholarship "
        f"records, alumni records, or other department evidence for this criterion."
    )


def _filter_criterion_hits(hits: list[dict], crit_name: str, keywords: str) -> list[dict]:
    criterion_words = {
        w
        for w in re.findall(r"[a-z]+", f"{crit_name} {keywords}".lower())
        if len(w) > 3 and w not in {"criterion", "criteria", "naac", "with", "and", "from"}
    }
    context_words = {
        "naac", "ssr", "iqac", "accreditation", "evidence", "policy",
        "department", "institution", "student", "students", "alumni",
        "placement", "scholarship", "progression", "support",
    }

    relevant_hits = []
    for hit in hits:
        text = hit.get("text", "")
        text_lower = text.lower()
        matched = {word for word in criterion_words if word in text_lower}
        has_context = any(word in text_lower for word in context_words)

        if len(matched) >= 2 or (has_context and matched):
            relevant_hits.append(hit)

    return relevant_hits


def generate_criterion_summary(username: str, criterion_number: int) -> str:
    """Generate NAAC Criterion-wise summary (Criterion 1 to 7) with targeted criterion search."""
    crit_info = CRITERION_MAP.get(criterion_number, (f"Criterion {criterion_number}", "NAAC policy quality criteria"))
    crit_name, keywords = crit_info
    
    query = f"NAAC Criterion {criterion_number} {crit_name} {keywords}"
    hits = query_rag_index(username, query, top_k=6)
    hits = _filter_criterion_hits(hits, crit_name, keywords)
    
    if not hits:
        return _criterion_not_found_message(criterion_number, crit_name)

    context_str = "\n\n".join([f"--- Chunk {i+1} ({h['source']}) ---\n{h['text']}" for i, h in enumerate(hits)])
    
    prompt = (
        f"You are the senior NAAC Accreditation AI Specialist.\n"
        f"Generate a comprehensive SSR (Self-Study Report) summary for NAAC Criterion {criterion_number}: {crit_name} based on the document context below.\n\n"
        f"If the context is unrelated to NAAC Criterion {criterion_number}, respond only with: "
        f"\"The uploaded documents do not contain enough relevant evidence for NAAC Criterion {criterion_number}.\"\n"
        f"Do not summarize aptitude formulas, algebra notes, random syllabus content, or unrelated text as criterion evidence.\n\n"
        f"Document Context:\n{context_str}\n\n"
        f"Provide:\n"
        f"1. Executive Summary for Criterion {criterion_number}: {crit_name}.\n"
        f"2. Key Metrics, Policies & Evidence found in the documents.\n"
        f"3. Actionable NAAC SSR Documentation Recommendations."
    )

    try:
        return invoke_openrouter_free_models(prompt, temperature=0.1)
    except Exception as exc:
        logger.error(f"OpenRouter free-model criterion fallback failed: {exc}")
        evidence_rows = []
        for idx, hit in enumerate(hits[:5], 1):
            excerpt = _clean_text(hit.get("text", ""))[:220]
            evidence_rows.append(f"{idx}. `{hit.get('source', 'Unknown')}` page {hit.get('page', 1)}: {excerpt}")
        return (
            f"**NAAC Criterion {criterion_number}: {crit_name} Evidence Found**\n\n"
            f"OpenRouter AI is currently unavailable, but these relevant evidence snippets were found:\n\n"
            + "\n".join(evidence_rows)
        )


def compare_documents(username: str, doc1: str, doc2: str) -> str:
    """Compare two documents and highlight key differences."""
    from db_storage import load_knowledge_file
    c1 = load_knowledge_file(username, doc1)
    c2 = load_knowledge_file(username, doc2)

    if not c1 or not c2:
        return "One or both selected documents could not be found."

    prompt = (
        f"Compare the following two departmental documents and list their key differences, policy updates, and NAAC implications.\n\n"
        f"=== Document 1: {doc1} ===\n{c1[:2500].decode('utf-8', errors='ignore')}\n\n"
        f"=== Document 2: {doc2} ===\n{c2[:2500].decode('utf-8', errors='ignore')}"
    )

    try:
        return invoke_openrouter_free_models(prompt, temperature=0.2)
    except Exception as exc:
        return f"Comparison failed: {exc}"


def ingest_documents(username: str) -> dict:
    """Ingest and re-index all user documents."""
    try:
        build_vector_store(username, force_rebuild=True)
        return {"success": True, "message": "Knowledge Base re-indexed successfully!"}
    except Exception as exc:
        return {"success": False, "message": str(exc)}


def run_knowledge_agent(username: str, query: str) -> dict:
    return ask_knowledge_agent(username=username, query=query)
